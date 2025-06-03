import docx
from docx.shared import Pt
#import docx.styles
import pandas as pd
import sys 
import os 
from PyQt5 import QtWidgets, QtGui, QtCore
#from PyQt5.QtWidgets import QSizePolicy
from PyQt5.QtCore import *
from PyQt5.QtWidgets import QFileDialog, QTableWidgetItem
import mainform_new
import random
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
#from docx.enum import WD_LIST_NUMBERING
from PyQt5.QtGui import *
import math

class testgen(QtWidgets.QMainWindow, mainform_new.Ui_MainWindow):
   
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        icon_path = resource_path("logo.png")
        self.setWindowIcon(QtGui.QIcon(icon_path))
        self.action_load_file.triggered.connect(self.open_file_dialog)
        self.action_testgen.triggered.connect(self.gen)
        self.action_loadheader.triggered.connect(self.loadHeader)
        self.action_loadfooter.triggered.connect(self.loadFooter)
        self.actionupdate_list.triggered.connect(self.update_combolist)
        self.actionload_themes.triggered.connect(self.gen_themes)
        self.radioButton_1col.toggled.connect(lambda: self.on_radio_button_clicked(self.radioButton_1col))
        self.radioButton_2col.toggled.connect(lambda: self.on_radio_button_clicked(self.radioButton_2col))
        self.spinBox.valueChanged.connect(lambda: self.checkValue(self.spinBox, self.spinBox_2))
        self.spinBox_2.valueChanged.connect(lambda: self.checkValue(self.spinBox, self.spinBox_2))
        #self.form_layout = QtWidgets.QFormLayout()       
        self.question_count=0
        self.list_cmb = []
        self.themes = []
        self.header_content=''
        self.footer_content= ''
        self.text_col = 1
    #считывание .docx файла 
    def read_docx_file(self, file_path):
        document = docx.Document(file_path)
        text = []
        for paragraph in document.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    
    #загрузка шапки
    def loadHeader(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open File", "", "Word files (*.docx)")
        if file_path:
            try:
                self.header_content = self.read_docx_file(file_path)
                self.plainTextEdit_header.setPlainText(self.header_content)
            except Exception as e:
                self.dialog_critical(str(e))
    
    #загрузка футера
    def loadFooter(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open File", "", "Word Files (*.docx)")
        if file_path:
            try:
                self.footer_content = self.read_docx_file(file_path)
                self.plainTextEdit_footer.setPlainText(self.footer_content)
                
            except Exception as e:
                self.dialog_critical(str(e))
    
    def dialog_critical(self, s):
        dlg = QtWidgets.QMessageBox(self)
        dlg.setText(s)
        dlg.setIcon(QtWidgets.QMessageBox.Critical)
        dlg.show()

    #количество колонок в тесте
    def on_radio_button_clicked(self, rbtn):
        if rbtn.isChecked():
            self.text_col = int(rbtn.text())
    
    #доступность кнопки генерации
    def checkValue(self, spin1, spin2):
        if (spin1.value() == 0 or spin2.value() == 0):
            self.btn_testgen.setEnabled(False)
        else:
            self.btn_testgen.setEnabled(True)
            self.question_count=self.spinBox_2.value()
    
    #обновление  поля разбивки по темам        
    def update_combolist(self):
        #нужно очистить layout
        self.clear_form_layout()
        self.append_combos()
    
    def clear_form_layout(self):
        while self.formLayout.rowCount():
        # Получаем последнюю строку и её виджеты
            row_index = self.formLayout.rowCount() - 1
            label_item = self.formLayout.itemAt(row_index, QtWidgets.QFormLayout.LabelRole)
            field_item = self.formLayout.itemAt(row_index, QtWidgets.QFormLayout.FieldRole)

            if label_item is not None:
            # Убираем ссылку родителя у каждого виджета
                label_widget = label_item.widget()
                if label_widget is not None:
                    label_widget.deleteLater()

            if field_item is not None:
                field_widget = field_item.widget()
                if field_widget is not None:
                    field_widget.deleteLater()

        # Удаляем саму строку из макета
            self.formLayout.removeRow(row_index)
    
    #обновление тем в комбо
    def update_combo(self, index):#формирование списка тем из комбо
        sender_combo = self.sender()  # Получаем объект, вызвавший событие
        new_text = sender_combo.currentText()
        name = sender_combo.objectName()
        for i, cmb in enumerate(self.list_cmb):
            if cmb.objectName() == name:
                self.themes[i] = new_text
                break
    
    #добавление комбо в список
    def append_combos(self):
        self.list_cmb= []
        self.themes =[]
        for i in range(self.question_count):
            cmb = QtWidgets.QComboBox()
            sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Expanding)
            sizePolicy.setHeightForWidth(cmb.sizePolicy().hasHeightForWidth())
            cmb.setSizePolicy(sizePolicy)
            cmb.setMinimumSize(QtCore.QSize(300,70))
            cmb.setMaximumSize(QtCore.QSize(16777215, 70))
            font = QFont('Times', 12)
            font.setBold(False)
            font.setItalic(True)
            cmb.setFont(font)

            cmb.addItems(themes_set)
            cmb.setObjectName(f"cmb_{i}")
            self.list_cmb.append(cmb)
            self.themes.append(cmb.currentText())
            label = QtWidgets.QLabel()
            label.setFont(font)
            label.setText(f'Вопрос №{i+1}')
            
            self.formLayout.addRow(label, cmb)
            
        for combo in self.list_cmb:
            combo.currentIndexChanged.connect(self.update_combo)

    def gen_themes(self):
        #self.clear_form_layout()
        self.append_combos()
    
    #удаление пустых
    def del_nan(self, data):
         filtered_data = {k: v for k, v in data.items() if not (isinstance(v, float) and math.isnan(v))}
         return filtered_data
                   
    def open_file_dialog(self):
        try:
            file_path, _ = QFileDialog.getOpenFileName(self, "Open File", "", "Таблицы Excel (*.xlsx)")
            if file_path:
                df = pd.read_excel(file_path)
                global dict_list
                global themes_set 
                themes_list = []
                dict_list_temp=  df.to_dict(orient='records')
            
                dict_list = list(map(lambda x: self.convert_to_str(x), dict_list_temp))
           
                global keys_list
                keys_list=list(dict_list[0].keys())
            
                for i in range( len(dict_list)):
                    value =  dict_list[i].get(keys_list[0])
                #if type(value) == str:
                    themes_list.append(value)
                #else:
                #    values = str(value)
                #    themes_list.append(values)
                   
                themes_set  = set(themes_list) 
                   
                table = self.tableWidget
                global row_count
                global column_count
                row_count = (len(dict_list))
                column_count = (len(dict_list[0]))
                table.setColumnCount(column_count) 
                table.setRowCount(row_count)
           
                table.setHorizontalHeaderLabels((list(dict_list[0].keys())))
                       
                for row in range(row_count): 
                    for column in range(column_count):
                        table.setColumnWidth(column, 300)
                    
                        item = list(dict_list[row].values())[column]
                        print(item)
                        if ( pd.isna(item)):
                            item = ''
                        
                        table.setItem(row, column, QTableWidgetItem(item))
        except FileNotFoundError as e:            
            QtWidgets.QMessageBox.critical(self, 'Error', f'Ошибка чтения файла: {e}', QtWidgets.QMessageBox.Yes)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, 'Error', f'Ошибка: {e}', QtWidgets.QMessageBox.Yes)
    #добавление в тест шапки
    def add_header_section(self, doc, test, style):
        sectionH = doc.sections[0]
        
        sectionH.start_type = WD_SECTION.NEW_PAGE
        sectPr = sectionH._sectPr 
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '1')
        content = self.plainTextEdit_header.toPlainText()
        if self.headergroup.isChecked():
            if self.header_content.strip() !='' :
                head = doc.add_paragraph(content, style = style)
            else:
                 QtWidgets.QMessageBox.critical(self, 'Error', f'Поле заголовка пустое!', QtWidgets.QMessageBox.Yes)
                 return 
            head=doc.add_paragraph('Вариант '+str(test+1), style = style)

    #добавление в тест футера    
    def add_footer_section(self, doc):
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        sectionf = doc.add_section(WD_SECTION.CONTINUOUS)
        sectionf = doc.sections[-1]
        sectPr = sectionf._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '1')
        content= self.plainTextEdit_footer.toPlainText()
        if self.footergroup.isChecked():
            if content.strip()!= '':
                foot = doc.add_paragraph(content).style = style
            else:
                QtWidgets.QMessageBox.critical(self, 'Error', f'Поле футера пустое!', QtWidgets.QMessageBox.Yes)
                return
        doc.add_section(WD_SECTION.NEW_PAGE)

    #добавление таблицы ключей в файл    
    def add_keys_table(self, doc, question_count, list_keys):
        # добавляем таблицу с одной строкой для заполнения названий колонок
        table = doc.add_table(1, question_count+1)
        table.style = 'Table Grid'
        # Получаем строку с колонками из добавленной таблицы 
        head_cells = table.rows[0].cells
        #print(list_keys)
        # добавляем названия колонок
        p = head_cells[0].paragraphs[0]
        p.add_run(f'Вариант').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(1, question_count+1):
            #print(i)
            p = head_cells[i].paragraphs[0]
            p.add_run(f'Вопрос {i}').bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in range(len(list_keys)):
            cells = table.add_row().cells
            cells[0].text = f'Вариант {row+1}'
            for col in range(1,len(list_keys[row])+1):
                cells[col].text = str(list_keys[row][col-1])

    def convert_to_str(self, data):
        for key, value in data.items():
            if isinstance(value, (int, float)) and not math.isnan(value):
                data[key] = str(value)
        return data
    
    #генерация тестов
    def gen(self):
        def gen_ques(new_list):
            n = random.randint(1, len(new_list)-1)
            ques_text = doc.add_paragraph(f'Вопрос №{str(ques+1)}', style=style_main)
            ques_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
            item = list(new_list[n].values())
           # print(item)                                        
            doc.add_paragraph(item[1], style=style_main)
            for i in range(2,6):
                if (pd.isna(item[i]) != True):
                    par =doc.add_paragraph(f'{i-1}. '+str(item[i]), style=style_main)
                                                        
            list_ques.insert(ques, item[6])
                        
            doc.add_paragraph('________________________________')
        
        LEFT_INDENT = Pt(36)
        outfile_path, _ = QFileDialog.getSaveFileName(self, "Save File", "", "All Files (*)")
        try:      
        #if outfile_path:
            doc = docx.Document()
            # задаем стиль текста по умолчанию
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(14)
            #стиль для шапки
            styleH = doc.styles.add_style('MyHeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
            styleH.font.name = 'Arial'
            styleH.font.size = Pt(10)
            #styleH.font.underline = True
            styleH.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #_______________________________
            #стиль теста
            style_main = doc.styles.add_style('mainStyle', WD_STYLE_TYPE.PARAGRAPH)
            style_main.font.name = 'Arial'
            style_main.font.size = Pt(12)
            style_main.paragraph_format.algnment = WD_ALIGN_PARAGRAPH.JUSTIFY
            #_______________________________
            test_count = self.spinBox.value()
            
            self.question_count = self.spinBox_2.value()
            list_keys= []
            if (test_count != 0 and self.question_count != 0):
                
                for test in range(test_count):
                    self.add_header_section( doc, test, styleH)
                                   
                    doc.add_section(WD_SECTION.CONTINUOUS)
                    
                    section1 = doc.sections[-1]
                    sectPr = section1._sectPr
                    cols = sectPr.xpath('./w:cols')[0]
                    cols.set(qn('w:num'), str(self.text_col))
                    cols.set(qn('w:space'), '10')  

                    list_ques = []
                    ready_list = []
                    #вычленяем из словаря только тема-вопрос-ответы-номер_прав
                    # keys = ['Тема', 'Вопрос', 'Ответ1', 'Ответ2', 'Ответ3', 'Ответ4', 'Номер правильного']
                    #ready_list = []
                    for dict1 in dict_list:
                        dict2 = {}  # Создаем новый словарь на каждой итерации
                        for key in keys_list:
                            dict2[key] = dict1.get(key)
                            #Используем метод get(), чтобы избежать ошибок, если ключа нет
                        ready_list.append(dict2)
                    #print(ready_list)
                    
                    if not(self.themesgroupBox.isChecked()):
                        for ques in range(self.question_count):
                           gen_ques(self.ready_list) 
                    else:
                    #генерим по темам
                        for ques in range(self.question_count):
                            current_theme = self.themes[ques]
                            thema = keys_list[0]
                            #temp_list =list(map( lambda x: self.convert_to_str(x, thema), ready_list))
                            #print(temp_list)
                            # Проверка наличия нужной темы среди ключей
                            #new_list = list(filter(lambda x: x['Тема'] == current_theme, ready_list))
                            new_list = list(filter(lambda x: x[thema] == current_theme, ready_list))
                             
                                #print(new_list)
                            gen_ques(new_list)
                            
                    self.add_footer_section(doc)
                    #doc.add_section(WD_SECTION.NEW_PAGE)

                    list_keys.insert(test, list_ques)
                    #doc.add_page_break()
                # Добавление таблицы
                self.add_keys_table(doc, self.question_count, list_keys)
                
            elif (test_count == 0 or self.question_count == 0):
                self.btn_testgen.setEnabled(False)     
                                                        
            doc.save(outfile_path)
            doc.styles['MyHeaderStyle'].delete()
            doc.styles['mainStyle'].delete()
            
            QtWidgets.QMessageBox.information(self, 'Information', 'Файл сформирован', QtWidgets.QMessageBox.Yes)
        #else: 
        except OSError as err:
            QtWidgets.QMessageBox.critical(self, 'Error', f'Ошибка записи в файл:  {err}', QtWidgets.QMessageBox.Yes)

def resource_path(relative_path):
    #"""Получить абсолютный путь к ресурсу, работает для разработки и PyInstaller"""
    try:
        # PyInstaller создает временную папку и сохраняет путь в _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        # Если это не PyInstaller, используем текущую папку
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)
     
def main():
    from PyQt5.QtWinExtras import QtWin 
    myappid = 'mycompany.myproduct.subproduct.version'                         
    QtWin.setCurrentProcessExplicitAppUserModelID(myappid) 
    icon_path = resource_path("logo.png")
    app = QtWidgets.QApplication(sys.argv)
    #app.setWindowIcon(QtGui.QIcon(icon_path))
     
    window = testgen()  
    window.show() 
    app.exec_() 

if __name__ == '__main__':  
    main()  